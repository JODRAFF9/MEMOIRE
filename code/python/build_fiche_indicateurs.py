# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

doc = Document()
# base style
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5); st.font.color.rgb = BLACK
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

def set_cell_border(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for edge in ('top','bottom'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'), kw.get(edge,'4'))
        e.set(qn('w:color'),'000000'); tb.append(e)
    tcPr.append(tb)

def shade(cell, hexcol):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcol)
    tcPr.append(sh)

# Title
h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("Construction des indicateurs de pauvreté multidimensionnelle de l'enfant (N-MODA)")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLACK

# Principe (single-cell bordered box)
pt = doc.add_table(rows=1, cols=1); pt.alignment = WD_TABLE_ALIGNMENT.CENTER
c = pt.cell(0,0)
for edge in ('top','bottom','left','right'):
    tcPr = c._tc.get_or_add_tcPr(); tb = tcPr.find(qn('w:tcBorders'))
    if tb is None: tb = OxmlElement('w:tcBorders'); tcPr.append(tb)
    e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); tb.append(e)
p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
def add(p, txt, bold=False, italic=False):
    rr = p.add_run(txt); rr.bold = bold; rr.italic = italic; rr.font.size = Pt(9.5); rr.font.color.rgb = BLACK; return rr
add(p,"Principe. ",bold=True)
add(p,"Approche N-MODA (ANSD/UNICEF) : sept dimensions du bien-être, déclinées en indicateurs binaires (privé / non privé). Un enfant est ")
add(p,"privé dans une dimension",bold=True); add(p," s'il l'est dans ")
add(p,"au moins un",italic=True); add(p," de ses indicateurs (union intra-dimension). Il est ")
add(p,"pauvre (N-MODA)",bold=True); add(p," s'il est privé dans ")
add(p,"au moins k = 4 dimensions sur 7",bold=True)
add(p,". Unité d'analyse : l'enfant ; trois groupes d'âge (0-4, 5-14, 15-17 ans) déterminent les indicateurs applicables. "
      "Statistiques descriptives pondérées par le poids de sondage (hhweight) ; effets causals pondérés par les poids d'appariement du PSM.")
doc.add_paragraph()

DIMS = [
 ("1. Assainissement", [
   ("Type de sanitaire non amélioré","le ménage utilise des toilettes non améliorées : latrines SANPLAT, latrines dallées simples, fosse rudimentaire, toilettes publiques, aucune toilette, autre (codes 7 à 12).","s11q55 / s11q54"),
   ("Partage des toilettes","le ménage partage ses toilettes avec d'autres ménages (dénominateur : ménages à installation privée).","s11q56 / s11q55"),
 ]),
 ("2. Eau", [
   ("Source d'eau non améliorée","la source de boisson est non améliorée (puits ouvert cour/concession 5, puits ouvert ailleurs 6, source non aménagée 12, fleuve/rivière/lac 13, vendeur ambulant 16, autre 17, à l'une ou l'autre saison), et le ménage ne traite pas son eau (filtre de traitement différent de « oui »).","Source : s11q27a, s11q27b / s11q26a, s11q26b. Filtre traitement : s11q32 / s11q31"),
   ("Temps d'accès à l'eau","le temps de collecte (aller plus attente à la source) dépasse 30 min à l'une ou l'autre saison.","s11q29a, s11q31a (+ h/min) / s11q28a, s11q30a"),
 ]),
 ("3. Logement", [
   ("Débarras des ordures","le ménage évacue ses ordures de façon inadéquate : brûlées (3), dépotoir sauvage (5), autre (6).","s11q54 / s11q53"),
   ("Surpeuplement","au moins 4 personnes par pièce (ratio taille/pièces supérieur ou égal à 4 ; lecture ANSD de « plus de 3 personnes par pièce »).","hhsize (roster) / s11q02 (pièces)"),
 ]),
 ("4. Nutrition", [
   ("Insécurité alimentaire (FIES)","score FIES complet (8 questions) : inquiétude de manquer de nourriture, impossibilité de manger sainement, alimentation peu variée, saut d'un repas, avoir mangé moins que nécessaire, plus de nourriture, avoir eu faim sans manger, journée entière sans manger. Score de 0 à 8 ; privé si score >= 1. Manquants et 98/99 traités comme « non ». La diversité des repas n'est pas retenue : le module n'a pas été collecté en 2021.","s08aq01 à s08aq08"),
 ]),
 ("5. Santé  (dimension = union des deux indicateurs)", [
   ("Combustible solide pour cuisiner","le ménage cuisine avec un combustible solide : bois (ramassé/acheté), charbon de bois, déchets animaux, autres.","s11q53__1,2,3,7,8 / s11q52"),
   ("Accès à une structure de santé","l'enfant ne peut accéder à pied à une structure de santé publique (hôpital, service 5 ; autre centre de santé public, service 6) : le mode habituel pour rejoindre la plus proche n'est pas la marche.","Module communautaire s02_co : s02q00, s02q01__5/__6, s02q02"),
 ]),
 ("6. Protection de l'enfant", [
   ("Absence d'acte de naissance","l'enfant (moins de 15 ans) ne dispose pas d'un acte de naissance.","s01q05"),
   ("Travail des enfants","l'enfant (5-14 ans) effectue un travail économique ou domestique d'au moins 1 heure. Domestique : heures de courses, travaux domestiques, garde, corvées d'eau et de bois. Économique : travail rémunéré ou pour propre compte.","Domestique : s04q01 à s04q05. Économique : s04q06 à s04q09"),
   ("Séparation parentale","l'enfant ne vit pas avec ses deux parents biologiques (au moins un parent n'habite pas le ménage).","s01q22 (père) et s01q29 (mère)"),
 ]),
 ("7. Éducation", [
   ("Non-scolarisation","l'enfant (5-14 ans) n'est pas scolarisé.","scol (base individus)"),
   ("Illettrisme","l'enfant (15-17 ans) ne sait ni lire ni écrire.","alfab / alfa (base individus)"),
   ("NEET","Indicateur écarté. La valeur ANSD (85,7 %) correspond à la seule absence d'emploi, sans la condition de non-scolarisation, et compte donc les élèves comme privés : inutilisable comme privation éducative.","(scol, activ7j)"),
 ]),
]

for title, rows in DIMS:
    hp = doc.add_heading(level=2); hp.paragraph_format.space_before = Pt(8); hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run(title); hr.bold = True; hr.font.size = Pt(12.5); hr.font.color.rgb = BLACK
    t = doc.add_table(rows=1, cols=3); t.autofit = False
    widths = (Cm(3.6), Cm(9.8), Cm(4.2))
    hdr = t.rows[0].cells
    for i,htxt in enumerate(("Indicateur","Définition (privé si...)","Variables (2018 / 2021)")):
        hdr[i].width = widths[i]; set_cell_border(hdr[i], top='12', bottom='12')
        pp = hdr[i].paragraphs[0]; rr = pp.add_run(htxt); rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = BLACK
    for (ind,defn,var) in rows:
        cells = t.add_row().cells
        for i,val in enumerate((ind,defn,var)):
            cells[i].width = widths[i]; set_cell_border(cells[i], bottom='4')
            pp = cells[i].paragraphs[0]; pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(val); rr.font.size = Pt(9.5); rr.font.color.rgb = BLACK
            if i == 0: rr.bold = True

fp = doc.add_paragraph(); fp.paragraph_format.space_before = Pt(10)
fr = fp.add_run("Note : les variables 2021 sont précisées après le « / » lorsqu'elles diffèrent de 2018. "
   "Les âges en gras restreignent l'indicateur au groupe concerné ; ailleurs l'indicateur, mesuré au niveau du "
   "ménage ou de la localité, s'applique à tous les groupes d'âge.")
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = BLACK

out = "/home/user/MEMOIRE/Docs/construction_indicateurs_nmoda.docx"
doc.save(out)
print("saved", out)
